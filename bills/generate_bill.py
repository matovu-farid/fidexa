"""Generate a water utility bill (.docx) modelled on NWSC tax invoices."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HEADER_FILL = "BDD6EE"          # light blue header band
HEADER_TEXT = RGBColor(0x1F, 0x37, 0x63)
ACCENT_BLUE = RGBColor(0x1F, 0x37, 0x63)
NOTICE_FILL = "BDD6EE"
TOTAL_FILL = "9CC2E5"
BORDER_COLOR = "000000"


@dataclass
class Bill:
    # Issuer
    company_name: str = "NATIONAL WATER & SEWERAGE CORPORATION"
    company_address: str = "P.O BOX 7053, PLOT 39 JINJA RD, KAMPALA"
    company_toll: str = "TOLL FREE: 0800200977 / 0800300977"

    # Customer block (text, top-right of header)
    customer_name: str = "FARID MATOVU NKOBA"
    customer_tel: str = "0705222144"
    customer_email: str = "matovu90@gmail.com"
    invoice_number: str = "2144820318555143-F"
    fdn: str = "22376085513126184277"
    etax_tin: str = "1000023440"

    # Invoice / customer reference
    invoice_date: date = date(2026, 5, 15)
    customer_reference: str = "21448291"

    # Service location
    service_address: str = "Sir Apollo Kaggwa Road, Kampala, Central Region 256, UG"
    property_reference: str = "21/35/18/612"

    # Basis & meter
    basis_of_charge: str = "Domestic ( Domestic Metered )"
    meter_details: str = "22-NIN15-104782"

    # Meter readings
    current_reading_date: date = date(2026, 5, 14)
    current_reading: int = 461
    previous_reading_date: date = date(2026, 4, 15)
    previous_reading: int = 428

    # Charges
    balance_bfwd: int = 0
    payments_since_prev: int = 0
    adjustments_since_prev: int = 0
    water_rate: int = 4_141          # UGX per m³
    include_sewerage: bool = False
    sewerage_pct: int = 75
    service_charge: int = 1_500
    vat_rate: float = 0.18

    @property
    def consumption(self) -> int:
        return self.current_reading - self.previous_reading

    @property
    def bfwd_as_at(self) -> int:
        return self.balance_bfwd + self.payments_since_prev + self.adjustments_since_prev

    @property
    def water_charge(self) -> int:
        return self.consumption * self.water_rate

    @property
    def sewerage_charge(self) -> int:
        if not self.include_sewerage:
            return 0
        return round(self.water_charge * self.sewerage_pct / 100)

    @property
    def pre_vat(self) -> int:
        return self.water_charge + self.sewerage_charge + self.service_charge

    @property
    def vat(self) -> int:
        return round(self.pre_vat * self.vat_rate)

    @property
    def sub_total(self) -> int:
        return self.pre_vat + self.vat

    @property
    def total_amount_due(self) -> int:
        return self.sub_total + self.bfwd_as_at


# ---------- formatting helpers ----------

def fmt_money(n: int) -> str:
    return f"{n:,}" if n >= 0 else f"-{abs(n):,}"


def fmt_date(d: date) -> str:
    return d.strftime("%d/%m/%Y")


# ---------- low-level XML helpers ----------

def shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = BORDER_COLOR, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def set_table_borders(table, color: str = BORDER_COLOR, size: str = "4") -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color, size)


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    size: int = 10,
    align: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    r.font.size = Pt(size)


def header_cell(cell, text: str, *, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    shade_cell(cell, HEADER_FILL)
    write_cell(cell, text, bold=True, color=HEADER_TEXT, size=10, align=align)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def data_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    size: int = 10,
    fill: str | None = None,
) -> None:
    if fill:
        shade_cell(cell, fill)
    write_cell(cell, text, bold=bold, size=size, align=align)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def merge_horizontal(row, start: int, end: int):
    cell = row.cells[start]
    for i in range(start + 1, end + 1):
        cell = cell.merge(row.cells[i])
    return cell


# ---------- watermark ----------

_SHAPETYPE_XML = (
    '<v:shapetype xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" '
    'path="m@7,l@8,m@5,21600l@6,21600e">'
    "<v:formulas>"
    '<v:f eqn="sum #0 0 10800"/>'
    '<v:f eqn="prod #0 2 1"/>'
    '<v:f eqn="sum 21600 0 @1"/>'
    '<v:f eqn="sum 0 0 @2"/>'
    '<v:f eqn="sum 21600 0 @3"/>'
    '<v:f eqn="if @0 @3 0"/>'
    '<v:f eqn="if @0 21600 @1"/>'
    '<v:f eqn="if @0 0 @2"/>'
    '<v:f eqn="if @0 @4 21600"/>'
    '<v:f eqn="mid @5 @6"/>'
    '<v:f eqn="mid @8 @5"/>'
    '<v:f eqn="mid @7 @8"/>'
    '<v:f eqn="mid @6 @7"/>'
    '<v:f eqn="sum @6 0 @5"/>'
    "</v:formulas>"
    '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="custom" '
    'o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" '
    'o:connectangles="270,180,90,0"/>'
    '<v:textpath on="t" fitshape="t"/>'
    "<v:handles>"
    '<v:h position="#0,bottomRight" xrange="6629,14971"/>'
    "</v:handles>"
    '<o:lock v:ext="edit" text="t" shapetype="t"/>'
    "</v:shapetype>"
)


def add_watermark(doc: Document, text: str) -> None:
    """Insert a diagonal text watermark on every page (via header VML)."""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    shape_xml = (
        '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'id="WatermarkShape" o:spid="_x0000_s2049" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;'
        "width:520pt;height:160pt;rotation:315;z-index:-251654144;"
        "mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        'mso-position-vertical-relative:margin" '
        'o:allowincell="f" fillcolor="#D9E2F3" stroked="f">'
        '<v:textpath style="font-family:&apos;Calibri&apos;;font-size:1pt;'
        'v-text-kern:t" trim="t" fitpath="t" string="' + text + '"/>'
        "</v:shape>"
    )

    pict = OxmlElement("w:pict")
    pict.append(parse_xml(_SHAPETYPE_XML))
    pict.append(parse_xml(shape_xml))
    run._r.append(pict)


_DRAWINGML_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def add_image_watermark(
    doc: Document,
    image_path: Path,
    width_pt: int = 460,
    height_pt: int = 380,
) -> None:
    """Insert a centered, rotated image watermark on every page (via header VML)."""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    run.add_picture(str(image_path), width=Cm(4))
    blip = run._r.find(f".//{{{_DRAWINGML_NS}}}blip")
    rid = blip.get(f"{{{_RELATIONSHIPS_NS}}}embed")

    drawing = run._r.find(qn("w:drawing"))
    run._r.remove(drawing)

    shape_xml = (
        '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'id="WatermarkImage" o:spid="_x0000_s2050" type="#_x0000_t75" '
        'style="position:absolute;margin-left:0;margin-top:0;'
        f"width:{width_pt}pt;height:{height_pt}pt;rotation:315;z-index:-251654144;"
        "mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        'mso-position-vertical-relative:margin" '
        'o:allowincell="f">'
        f'<v:imagedata r:id="{rid}" o:title="watermark"/>'
        "</v:shape>"
    )

    pict = OxmlElement("w:pict")
    pict.append(parse_xml(shape_xml))
    run._r.append(pict)


# ---------- document sections ----------

def add_company_header(doc: Document, bill: Bill, logo_path: Path) -> None:
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(2)
    logo_p.add_run().add_picture(str(logo_path), width=Cm(3.2))

    for text, size, bold in (
        (bill.company_address, 10, False),
        (bill.company_toll, 10, False),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)


def add_customer_contact_block(doc: Document, bill: Bill) -> None:
    fields = [
        ("Name:", bill.customer_name),
        ("Tel:", bill.customer_tel),
        ("Email:", bill.customer_email),
        ("Invoice Number:", bill.invoice_number),
        ("FDN:", f"{bill.fdn}   eTAX TIN: {bill.etax_tin}"),
    ]
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{label} ")
        r.font.size = Pt(10)
        r.bold = True
        r = p.add_run(value)
        r.font.size = Pt(10)


def add_invoice_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("TAX INVOICE FOR WATER & SEWERAGE SERVICES")
    r.bold = True
    r.font.color.rgb = ACCENT_BLUE
    r.font.size = Pt(12)


def add_party_table(doc: Document, bill: Bill) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [8.0, 8.0])

    header_cell(table.rows[0].cells[0], "TO")
    header_cell(table.rows[0].cells[1], "INVOICE DATE")

    data_cell(table.rows[1].cells[0], bill.customer_name)
    data_cell(table.rows[1].cells[1], fmt_date(bill.invoice_date))

    data_cell(table.rows[2].cells[0], "")
    header_cell(table.rows[2].cells[1], "CUSTOMER REFERENCE")

    data_cell(table.rows[3].cells[0], "")
    data_cell(table.rows[3].cells[1], bill.customer_reference)

    set_table_borders(table)


def add_service_table(doc: Document, bill: Bill) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [8.0, 8.0])

    merged = merge_horizontal(table.rows[0], 0, 1)
    header_cell(merged, "IN RESPECT OF SERVICES AT")

    data_cell(table.rows[1].cells[0], bill.service_address)
    header_cell(table.rows[1].cells[1], "PROPERTY REFERENCE")

    data_cell(table.rows[2].cells[0], "")
    data_cell(table.rows[2].cells[1], bill.property_reference)

    set_table_borders(table)


def add_basis_meter_table(doc: Document, bill: Bill) -> None:
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [8.0, 8.0])

    merged = merge_horizontal(table.rows[0], 0, 1)
    header_cell(merged, "BASIS OF CHARGE")

    data_cell(table.rows[1].cells[0], bill.basis_of_charge)
    data_cell(table.rows[1].cells[1], "")

    header_cell(table.rows[2].cells[0], "METER DETAILS")
    data_cell(table.rows[2].cells[1], bill.meter_details)

    set_table_borders(table)


def add_readings_table(doc: Document, bill: Bill) -> None:
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [5.3, 5.3, 5.4])

    data_cell(table.rows[0].cells[0], "")
    header_cell(table.rows[0].cells[1], "DATE READ")
    header_cell(table.rows[0].cells[2], "READING")

    data_cell(table.rows[1].cells[0], "CURRENT", bold=True)
    data_cell(table.rows[1].cells[1], fmt_date(bill.current_reading_date))
    data_cell(table.rows[1].cells[2], str(bill.current_reading))

    data_cell(table.rows[2].cells[0], "PREVIOUS", bold=True)
    data_cell(table.rows[2].cells[1], fmt_date(bill.previous_reading_date))
    data_cell(table.rows[2].cells[2], str(bill.previous_reading))

    data_cell(table.rows[3].cells[0], "Consumption", bold=True)
    data_cell(table.rows[3].cells[1], "")
    data_cell(table.rows[3].cells[2], str(bill.consumption))

    set_table_borders(table)


def add_charges_table(doc: Document, bill: Bill) -> None:
    rows_data: list[tuple[str, int, bool]] = [
        ("Balance B/Fwd from previous Invoice", bill.balance_bfwd, False),
        ("Payments since Prev Invoice", bill.payments_since_prev, False),
        ("Adjustments since Prev Invoice", bill.adjustments_since_prev, False),
        (f"B/Fwd as at {fmt_date(bill.invoice_date)}", bill.bfwd_as_at, False),
        (
            f"WATER {bill.consumption} m3 at {bill.water_rate:,}/=",
            bill.water_charge,
            False,
        ),
        (
            f"SEWERAGE at {bill.sewerage_pct}% WATER CHARGES",
            bill.sewerage_charge,
            False,
        ),
        ("Service Charge", bill.service_charge, False),
        ("VAT", bill.vat, False),
        ("Sub Total", bill.sub_total, True),
        ("Total Amount Due", bill.total_amount_due, True),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [10.6, 5.4])

    header_cell(table.rows[0].cells[0], "CHARGING DETAILS")
    header_cell(table.rows[0].cells[1], "AMOUNT UGX", align=WD_ALIGN_PARAGRAPH.RIGHT)

    for i, (desc, amount, is_summary) in enumerate(rows_data, start=1):
        is_total = desc == "Total Amount Due"
        fill = TOTAL_FILL if is_total else None
        data_cell(
            table.rows[i].cells[0],
            desc,
            bold=is_summary,
            fill=fill,
        )
        data_cell(
            table.rows[i].cells[1],
            fmt_money(amount),
            bold=is_summary,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            fill=fill,
        )

    set_table_borders(table)


def add_notice_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_column_widths(table, [16.0])
    cell = table.rows[0].cells[0]
    shade_cell(cell, NOTICE_FILL)

    lines = [
        "Please pay your bill to avoid interruption of services.",
        "Clean your tank every month.",
        "Report any suspected unauthorized usage of water to our confidential line.",
    ]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = HEADER_TEXT

    set_table_borders(table)


# ---------- entry ----------

def build_document(bill: Bill, output_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    here = Path(__file__).parent
    add_image_watermark(doc, here / "logo-watermark.png")

    add_company_header(doc, bill, here / "logo.png")
    add_customer_contact_block(doc, bill)
    add_invoice_title(doc)

    add_party_table(doc, bill)
    add_service_table(doc, bill)
    add_basis_meter_table(doc, bill)
    add_readings_table(doc, bill)
    add_charges_table(doc, bill)
    add_notice_block(doc)

    doc.save(output_path)


def main() -> None:
    bill = Bill()
    output = Path(__file__).parent / "water-bill.docx"
    build_document(bill, output)
    print(f"Generated: {output}")
    print(f"Total amount due: UGX {fmt_money(bill.total_amount_due)}")


if __name__ == "__main__":
    main()
